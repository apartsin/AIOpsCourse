# -*- coding: utf-8 -*-
"""Build the HIT course-catalogue package for Modern AI Systems by cloning the
DLCourseHIT reference .docx templates (preserving the official HIT letterhead
header, styles, and fonts) and replacing the body with this course's content.
Outputs: hit-catalogue/{syllabus_en,syllabus_he,rationale,catalogue_summary}.docx
"""
import os
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
REF = os.path.join(ROOT, "build", "templates")
OUT = os.path.join(ROOT, "hit-catalogue")
os.makedirs(OUT, exist_ok=True)

# ---------- low-level helpers ----------
def clear_body(doc):
    body = doc.element.body
    for child in list(body):
        if child.tag == qn('w:sectPr'):
            continue
        body.remove(child)

def _set_rtl(p, rtl):
    pPr = p._p.get_or_add_pPr()
    if rtl:
        if pPr.find(qn('w:bidi')) is None:
            pPr.append(OxmlElement('w:bidi'))
        for run in p.runs:
            rPr = run._r.get_or_add_rPr()
            if rPr.find(qn('w:rtl')) is None:
                rPr.append(OxmlElement('w:rtl'))

def para(doc, text, style=None, rtl=False, align=None, bold=False, size=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    run = p.add_run(text)
    if bold: run.bold = True
    if size: run.font.size = Pt(size)
    if align is not None: p.alignment = align
    _set_rtl(p, rtl)
    return p

def heading(doc, text, rtl=False):
    p = para(doc, text, style='Heading 1', rtl=rtl,
             align=WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT)
    return p

def weekly_table(doc, header, rows, rtl=False):
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    if rtl:
        tblPr = t._tbl.tblPr
        if tblPr.find(qn('w:bidiVisual')) is None:
            tblPr.append(OxmlElement('w:bidiVisual'))
    hc = t.rows[0].cells
    for i, htxt in enumerate(header):
        hc[i].text = ""
        pr = hc[i].paragraphs[0]
        r = pr.add_run(htxt); r.bold = True
        if rtl: pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT; _set_rtl(pr, True)
    for wk, subj in rows:
        c = t.add_row().cells
        c[0].text = wk
        c[1].text = subj
        for cell in c:
            pp = cell.paragraphs[0]
            if rtl: pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT; _set_rtl(pp, True)
    return t

# ---------- shared content ----------
TITLE_EN = "AI Systems Engineering"
TITLE_HE = "הנדסת מערכות בינה מלאכותית"

BIB = [
    "Beyer, Betsy, Chris Jones, Jennifer Petoff, and Niall Richard Murphy, eds. Site Reliability Engineering: How Google Runs Production Systems. O'Reilly Media, 2016.",
    "Kleppmann, Martin. Designing Data-Intensive Applications. O'Reilly Media, 2017.",
    "Kim, Gene, Jez Humble, Patrick Debois, John Willis, and Nicole Forsgren. The DevOps Handbook. 2nd ed. IT Revolution Press, 2021.",
    "Reis, Joe, and Matt Housley. Fundamentals of Data Engineering. O'Reilly Media, 2022.",
    "Huyen, Chip. Designing Machine Learning Systems. O'Reilly Media, 2022.",
    "Huyen, Chip. AI Engineering: Building Applications with Foundation Models. O'Reilly Media, 2025.",
]

WEEKS_EN = [
    ("1", "Production engineering and the operations landscape; SLOs and error budgets; the five layers."),
    ("2", "Cloud computing fundamentals: compute, storage, and networking; deployment models (IaaS to serverless); cost."),
    ("3", "CI/CD, the testing pyramid, and versioned REST services; infrastructure as code."),
    ("4", "Orchestration with Kubernetes; deployment patterns (blue-green, canary); observability (logs, metrics, traces)."),
    ("5", "Data lakes and the lakehouse; the medallion architecture; pipelines, versioning, and lineage."),
    ("6", "Data quality and contracts; streaming with Kafka; feature stores and train/serve skew."),
    ("7", "Experiment tracking; the model registry and catalogue; model serving and safe rollout."),
    ("8", "Monitoring; data and concept drift; retraining triggers and model governance."),
    ("9", "LLM foundations: AI APIs, tokens and the token economy; structured outputs; managed AI platforms."),
    ("10", "Retrieval-augmented generation; vector databases; prompts as code; serving and the gateway pattern."),
    ("11", "LLM evaluation and eval sets; LLM-as-judge; observability and tracing; guardrails and prompt injection."),
    ("12", "Agents and AgentOps: tool use and function calling; the Model Context Protocol; tracing, bounds, and managed agents."),
    ("13", "Security and the supply chain; the OWASP Top 10 for LLM applications; governance and synthesis."),
]
WEEKS_HE = [
    ("1", "הנדסת ייצור ונוף התפעול; יעדי רמת שירות (SLO) ותקציבי שגיאה; חמש שכבות התפעול."),
    ("2", "יסודות מחשוב הענן: חישוב, אחסון ורשת; מודלי פריסה (מ-IaaS ועד serverless); עלות."),
    ("3", "CI/CD, פירמידת הבדיקות ושירותי REST מגורסים; תשתית כקוד."),
    ("4", "תזמור עם Kubernetes; דפוסי פריסה (blue-green, canary); נצפות (לוגים, מדדים, עקבות)."),
    ("5", "אגמי נתונים ו-lakehouse; ארכיטקטורת המדליון; צינורות, גרסאות ושושלת נתונים."),
    ("6", "איכות נתונים וחוזי נתונים; הזרמה עם Kafka; מאגרי מאפיינים והטיית אימון-הגשה."),
    ("7", "מעקב ניסויים; מרשם וקטלוג מודלים; הגשת מודלים והשקה בטוחה."),
    ("8", "ניטור; היסט נתונים והיסט מושגי; טריגרים לאימון מחדש וממשל מודלים."),
    ("9", "יסודות מודלי שפה: ממשקי API, טוקנים וכלכלת הטוקנים; פלטים מובנים; פלטפורמות AI מנוהלות."),
    ("10", "ייצור מבוסס אחזור (RAG); מסדי נתונים וקטוריים; פרומפטים כקוד; הגשה ודפוס שער (gateway)."),
    ("11", "הערכת מודלי שפה ומערכי הערכה; מודל-כשופט; נצפות ועקבות; מגיני בטיחות והזרקת פרומפט."),
    ("12", "סוכנים ו-AgentOps: שימוש בכלים וקריאות פונקציה; פרוטוקול MCP; עקבות, חסמים וסוכנים מנוהלים."),
    ("13", "אבטחה ושרשרת האספקה; OWASP Top 10 ליישומי מודלי שפה; ממשל וסינתזה."),
]

# ============================================================
# 1) syllabus_en.docx  (English, on HIT letterhead; LTR body)
# ============================================================
def build_syllabus_en():
    d = Document(os.path.join(REF, "syllabus_en.docx"))
    clear_body(d)
    L = WD_ALIGN_PARAGRAPH.LEFT
    para(d, TITLE_EN, bold=True, size=15, align=L)
    para(d, "DevOps, DataOps, MLOps, LLMOps, AgentOps", align=L)
    para(d, "Lecture: 2 hours, practice: 2 hours", align=L)
    para(d, "4 hours, 4 credits", align=L)
    para(d, "Prerequisites: Introduction to Machine Learning, Operating Systems, Software Engineering", align=L)

    heading(d, "Course Objectives")
    para(d, "Modern AI systems fail far more often in operations than in modelling. A model that scores well offline still has to be packaged, served, version-controlled, fed with trustworthy data, monitored for drift, secured, costed, and governed once real users depend on it. This course teaches the engineering discipline that surrounds the model: the practices, tooling, and architectures that take a prototype into reliable, observable, continuously-improving production service.", align=L)
    para(d, "It is an advanced course taken after a first machine-learning course. No prior cloud or large-language-model experience is assumed: the cloud foundations are built in week 2 and the LLM and AI-API foundations in week 9, before the course operates on either. By the end, students can take an AI-enabled service from specification through to a monitored, governed production deployment.", align=L)

    heading(d, "Course content")
    para(d, "The course is organised as a layered stack of operational practices. It begins with foundations and the cloud (compute, storage, networking, deployment models), then DevOps (CI/CD, testing, REST services, container orchestration with Kubernetes, and observability). It continues with DataOps (data lakes and the medallion architecture, pipelines and versioning, data quality and contracts, streaming, and feature stores) and MLOps (experiment tracking, the model registry and catalogue, serving and safe rollout, monitoring and drift). The final third covers LLMOps (AI APIs and the token economy, retrieval-augmented generation, gateways, evaluation, observability, and guardrails) and AgentOps (tool use, the Model Context Protocol, tracing, and bounded autonomous agents), closing with security and governance. It is highly practical and designed for working with an AI coding assistant.", align=L)

    heading(d, "Student duties and grade components")
    para(d, "The course is project-based with no written exams and no separate weekly labs. Teams of three or four carry a single AI-enabled service through the entire stack as one end-to-end system, built up in weekly increments. The final grade is the running project, graded at three Student Project Presentations: Specification (20%, week 5), Interim (30%, week 8), and Final with a short oral defense (50%, week 13).", align=L)

    heading(d, "Course of lessons")
    para(d, "Each week has two parts: a 2-hour lecture that develops the concepts and architecture, and a 2-hour practice session in which the instructor demonstrates the tooling live and teaches the hands-on topics that belong at the keyboard. Teaching methods: frontal teaching with slides and live code demonstrations. Use of technology: hands-on work in Python with industry-standard tooling (Docker, Kubernetes, MLflow, vector databases, LLM APIs and gateways) and an AI coding assistant. Guest lecturers: none.", align=L)
    para(d, "The order of the lessons (may change if required):", align=L)
    weekly_table(d, ("Week", "Subject"), WEEKS_EN, rtl=False)

    heading(d, "Textbooks")
    for b in BIB:
        para(d, b, align=L)

    d.save(os.path.join(OUT, "syllabus_en.docx"))
    print("wrote syllabus_en.docx")

# ============================================================
# 2) syllabus_he.docx  (Hebrew, RTL)
# ============================================================
def build_syllabus_he():
    d = Document(os.path.join(REF, "syllabus_he.docx"))
    clear_body(d)
    R = WD_ALIGN_PARAGRAPH.RIGHT
    para(d, TITLE_HE + " - " + TITLE_EN, rtl=True, align=R, bold=True, size=14)
    para(d, "DevOps, DataOps, MLOps, LLMOps, AgentOps", align=WD_ALIGN_PARAGRAPH.LEFT)
    para(d, "אופן הוראה: שיעור ותרגול.", rtl=True, align=R)
    para(d, "שעות שבועיות: הרצאה 2 שעות + תרגול 2 שעות, סה\"כ שעות – 4", rtl=True, align=R)
    para(d, "נקודות זכות: 4", rtl=True, align=R)
    para(d, "דרישות קדם: מבוא ללמידת מכונה, מערכות הפעלה, הנדסת תוכנה", rtl=True, align=R)

    heading(d, "מטרות הקורס", rtl=True)
    para(d, "מערכות בינה מלאכותית מודרניות נכשלות בתפעול הרבה יותר מאשר במידול. מודל שמשיג ביצועים טובים במעבדה עדיין צריך להיארז, להיות מוגש, מנוהל גרסאות, מוזן בנתונים אמינים, מנוטר להיסט, מאובטח, מתומחר וממושל ברגע שמשתמשים אמיתיים תלויים בו. קורס זה מלמד את הדיסציפלינה ההנדסית שמקיפה את המודל: הפרקטיקות, הכלים והארכיטקטורות שמעבירים אב-טיפוס לשירות ייצור אמין, נצפה ומשתפר בהתמדה.", rtl=True, align=R)
    para(d, "זהו קורס מתקדם הנלמד לאחר קורס מבוא בלמידת מכונה. לא נדרש ניסיון קודם בענן או במודלי שפה גדולים: יסודות הענן נבנים בשבוע 2 ויסודות מודלי השפה וממשקי ה-API בשבוע 9, לפני שהקורס מתפעל מי מהם. בסיום הקורס יוכלו הסטודנטים לקחת שירות מבוסס בינה מלאכותית מאפיון ועד פריסת ייצור מנוטרת וממושלת.", rtl=True, align=R)

    heading(d, "תוכן הקורס", rtl=True)
    para(d, "הקורס בנוי כמחסנית שכבתית של פרקטיקות תפעול. הוא נפתח ביסודות ובענן (חישוב, אחסון, רשת ומודלי פריסה), וממשיך ל-DevOps (CI/CD, בדיקות, שירותי REST, תזמור מכולות עם Kubernetes ונצפות). לאחר מכן DataOps (אגמי נתונים וארכיטקטורת המדליון, צינורות וגרסאות, איכות וחוזי נתונים, הזרמה ומאגרי מאפיינים) ו-MLOps (מעקב ניסויים, מרשם וקטלוג מודלים, הגשה והשקה בטוחה, ניטור והיסט). השליש האחרון עוסק ב-LLMOps (ממשקי API וכלכלת הטוקנים, ייצור מבוסס אחזור, שערים, הערכה, נצפות ומגיני בטיחות) וב-AgentOps (שימוש בכלים, פרוטוקול MCP, עקבות וסוכנים אוטונומיים חסומים), ונחתם באבטחה וממשל. הקורס מעשי מאוד ומותאם לעבודה עם עוזר תכנות מבוסס בינה מלאכותית.", rtl=True, align=R)

    heading(d, "חובות התלמידים ומרכיבי הציון", rtl=True)
    para(d, "הקורס מבוסס פרויקט, ללא מבחן כתוב וללא מעבדות שבועיות נפרדות. צוותים של שלושה עד ארבעה סטודנטים מעבירים שירות יחיד מבוסס בינה מלאכותית לאורך כל המחסנית כמערכת אחת מקצה לקצה, הנבנית בתוספות שבועיות. הציון הסופי הוא הפרויקט הרץ, המוערך בשלוש מצגות פרויקט סטודנטים: אפיון (20%, שבוע 5), ביניים (30%, שבוע 8), וסיום עם הגנה קצרה בעל-פה (50%, שבוע 13).", rtl=True, align=R)

    heading(d, "מהלך השיעורים", rtl=True)
    para(d, "כל שבוע כולל שני חלקים: הרצאה בת שעתיים המפתחת את המושגים והארכיטקטורה, ושיעור תרגול בן שעתיים שבו המרצה מדגים את הכלים בזמן אמת ומלמד את הנושאים המעשיים שמקומם במקלדת.", rtl=True, align=R)
    para(d, "שיטות ההוראה: הוראה פרונטלית מלווה במצגות ובהדגמות קוד.", rtl=True, align=R)
    para(d, "שימוש בטכנולוגיה: עבודה מעשית ב-Python עם כלים תעשייתיים (Docker, Kubernetes, MLflow, מסדי נתונים וקטוריים, ממשקי API ושערי מודלי שפה) ועוזר תכנות מבוסס בינה מלאכותית.", rtl=True, align=R)
    para(d, "מרצים אורחים: אין.", rtl=True, align=R)
    para(d, "תכנית הוראה מפורטת לכל השיעורים (סדר השיעורים צפוי להשתנות):", rtl=True, align=R)
    weekly_table(d, ("שבוע", "נושאים"), WEEKS_HE, rtl=True)

    heading(d, "ביבליוגרפיה", rtl=True)
    for b in BIB:
        para(d, b, align=WD_ALIGN_PARAGRAPH.LEFT)  # English citations left-aligned

    d.save(os.path.join(OUT, "syllabus_he.docx"))
    print("wrote syllabus_he.docx")

# ============================================================
# 3) rationale.docx  (Hebrew)
# ============================================================
def build_rationale():
    d = Document(os.path.join(REF, "rationale.docx"))
    clear_body(d)
    R = WD_ALIGN_PARAGRAPH.RIGHT
    para(d, "מסמך רציונל לקורס הנדסת מערכות בינה מלאכותית", rtl=True, align=R, bold=True, size=14)
    para(d, TITLE_EN, align=WD_ALIGN_PARAGRAPH.LEFT, bold=True)
    para(d, "הקורס עוסק בהנדסה, בפריסה, בניטור, בממשל ובתפעול של מערכות תוכנה מבוססות בינה מלאכותית בייצור. הוא מקיף חמש שכבות תפעול: DevOps, DataOps, MLOps, LLMOps ו-AgentOps, ומקנה בסיס מעשי לבניית מערכת מקצה לקצה הכוללת ענן, צינורות נתונים, הגשת מודלים, יכולת מבוססת מודל שפה וסוכן, ונצפות וממשל לאורך כולה.", rtl=True, align=R)
    para(d, "הקורס נדרש כקורס מתקדם בשנה ג' ומיועד לסטודנטים שסיימו קורס מבוא בלמידת מכונה ובעלי רקע בסיסי בהנדסת תוכנה ובמערכות הפעלה. הוא מספק את הבסיס המשותף לפריסה ולתפעול של בינה מלאכותית בעולם האמיתי, בסיס שכל התמחות בבינה מלאכותית (ראייה ממוחשבת, מודלי שפה, סוכנים) נפגשת בו בסופו של דבר. אין צורך בניסיון קודם בענן או במודלי שפה: הקורס בונה את שני היסודות הללו מאפס לפני שהוא מתפעל אותם.", rtl=True, align=R)
    para(d, "הקורס סוגר פער קיים בתכנית הלימודים: בעוד שקורסי הליבה מלמדים כיצד לבנות מודלים, קורס זה מלמד כיצד להריץ אותם בייצור באופן אמין, מאובטח, חסכוני וממושל. הוא מבוסס פרויקט ומותאם לאופן שבו הסטודנטים יעבדו בפועל, עם עוזר תכנות מבוסס בינה מלאכותית, תוך שמירה על למידה אמיתית: הסטודנט נדרש לבנות, לתפעל תחת עומס, ולהסביר ולהגן על כל החלטה הנדסית במצגות הפרויקט.", rtl=True, align=R)
    d.save(os.path.join(OUT, "rationale.docx"))
    print("wrote rationale.docx")

# ============================================================
# 4) catalogue_summary.docx  (Hebrew + English blurbs)
# ============================================================
def build_catalogue_summary():
    d = Document(os.path.join(REF, "catalogue_summary.docx"))
    clear_body(d)
    R = WD_ALIGN_PARAGRAPH.RIGHT; L = WD_ALIGN_PARAGRAPH.LEFT
    para(d, "תקצירים לידיעון", rtl=True, align=R, bold=True, size=14)
    para(d, TITLE_HE, rtl=True, align=R, bold=True)
    para(d, TITLE_EN, align=L)
    para(d, "אופן הוראה: שיעור ותרגול", rtl=True, align=R)
    para(d, "שעות שבועיות: הרצאה 2 שעות + תרגול 2 שעות, סה\"כ שעות – 4", rtl=True, align=R)
    para(d, "נקודות זכות: 4", rtl=True, align=R)
    para(d, "דרישות קדם: מבוא ללמידת מכונה, מערכות הפעלה, הנדסת תוכנה", rtl=True, align=R)
    para(d, "קורס מתקדם ומעשי בהנדסה, בפריסה ובתפעול של מערכות בינה מלאכותית בייצור, לאורך חמש שכבות: DevOps, DataOps, MLOps, LLMOps ו-AgentOps. הקורס בונה מאפס את יסודות הענן ואת יסודות מודלי השפה, ומשלב את החומר לכדי מערכת אחת מקצה לקצה הנבנית כפרויקט צוותי לאורך הסמסטר ומוצגת בשלוש מצגות. כולל עבודה עם עוזר תכנות מבוסס בינה מלאכותית.", rtl=True, align=R)
    para(d, "נושאי הקורס: ענן ו-DevOps (CI/CD, מכולות, Kubernetes, נצפות); DataOps (אגמי נתונים, ארכיטקטורת המדליון, איכות והזרמה); MLOps (מעקב, מרשם, הגשה, היסט); LLMOps (ממשקי API וכלכלת הטוקנים, RAG, שערים, הערכה ומגיני בטיחות); ו-AgentOps (כלים, MCP, עקבות וחסמים).", rtl=True, align=R)

    para(d, TITLE_EN, align=L, bold=True)
    para(d, "Lecture and practice", align=L)
    para(d, "4 hours, 4 credits", align=L)
    para(d, "Prerequisites: Introduction to Machine Learning, Operating Systems, Software Engineering", align=L)
    para(d, "An advanced, practical course on engineering, deploying, and operating AI systems in production across five layers: DevOps, DataOps, MLOps, LLMOps, and AgentOps. The course builds the cloud and the large-language-model foundations from scratch and integrates the material into one end-to-end system, built as a team project over the semester and presented three times. It includes work with an AI coding assistant.", align=L)
    para(d, "Topics: cloud and DevOps (CI/CD, containers, Kubernetes, observability); DataOps (data lakes, the medallion architecture, quality, and streaming); MLOps (tracking, registry, serving, drift); LLMOps (AI APIs and the token economy, RAG, gateways, evaluation, and guardrails); and AgentOps (tools, MCP, tracing, and bounds).", align=L)

    d.save(os.path.join(OUT, "catalogue_summary.docx"))
    print("wrote catalogue_summary.docx")

if __name__ == "__main__":
    build_syllabus_en()
    build_syllabus_he()
    build_rationale()
    build_catalogue_summary()
    print("done: 4 catalogue documents")
